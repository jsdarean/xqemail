from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

DOC_DIR = 'docs'
SCREENSHOT_DIR = os.path.join(DOC_DIR, 'screenshots')
OUTPUT_DOC = os.path.join(DOC_DIR, '需求催办系统_部署文档及操作手册.docx')

os.makedirs(DOC_DIR, exist_ok=True)

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
style.font.name = 'Microsoft YaHei'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
style.font.size = Pt(10.5)

# 标题样式
def add_heading_cn(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if level == 1:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(37, 99, 235)
    return heading

def add_paragraph_cn(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.bold = bold
    return p

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(31, 41, 55)
    return p

def add_screenshot(filename, caption):
    path = os.path.join(SCREENSHOT_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(6))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(f'图：{caption}')
        cap_run.font.size = Pt(9)
        cap_run.font.color.rgb = RGBColor(107, 114, 128)
        cap_run.font.name = 'Microsoft YaHei'
        cap_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[截图占位：{caption}]')
        run.font.color.rgb = RGBColor(239, 68, 68)

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run('注意：' + text)
    run.font.color.rgb = RGBColor(217, 119, 6)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ==================== 封面 ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('需求催办系统')
run.font.size = Pt(32)
run.font.bold = True
run.font.color.rgb = RGBColor(37, 99, 235)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('部署文档及操作手册')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(75, 85, 99)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run('版本：v1.4.0（可分发版）')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(107, 114, 128)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('日期：2026-07-01')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(107, 114, 128)
run.font.name = 'Microsoft YaHei'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()

# ==================== 目录 ====================
add_heading_cn('目录', level=1)
toc_items = [
    '1. 系统概述',
    '2. 环境要求',
    '3. 项目结构说明',
    '4. 部署步骤',
    '    4.1 解压项目',
    '    4.2 安装 Node.js 依赖',
    '    4.3 配置环境变量',
    '    4.4 初始化数据库',
    '    4.5 启动服务',
    '    4.6 验证部署',
    '    4.7 安装 Chrome 插件',
    '    4.8 启动插件本地 SMTP 中继服务器',
    '5. 日常使用操作手册',
    '    5.1 登录系统',
    '    5.2 查看需求清单',
    '    5.3 编辑工作量',
    '    5.4 发送邮件催办',
    '    5.5 发送微信催办',
    '    5.6 查看催办记录',
    '    5.7 配置 Web 端邮箱',
    '    5.8 退出登录',
    '    5.9 我的待办',
    '    5.10 集团需求导出',
    '    5.11 使用 Chrome 插件提取需求信息',
    '    5.12 插件收件人管理',
    '    5.13 插件邮箱与数据库设置',
    '6. 常见问题排查',
    '7. 安全建议',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.2 if not item.startswith('    ') else 0.6)

doc.add_page_break()

# ==================== 1. 系统概述 ====================
add_heading_cn('1. 系统概述', level=1)
add_paragraph_cn('需求催办系统是一款基于 Node.js + Express + MySQL 的 Web 应用，用于管理需求评估流程，支持以下核心功能：')
features = [
    'Web 端需求清单：按需求分组展示，支持展开/折叠、分页、排序',
    '工作量编辑：为每个需求填写工作量、标记是否涉及开发、填写开发单号，支持单行保存',
    '邮件催办：一键向未填写工作量的责任人发送催办邮件',
    '微信催办：自动生成催办文案并复制到剪贴板',
    '催办记录：查看历史催办记录及发送状态',
    'Web 端邮箱设置：可视化配置 SMTP/IMAP 邮箱参数',
    'Chrome 插件：在需求/工单页面一键提取信息并发送邮件，自动写入数据库',
    '我的待办：展示工作量已评估但开发单号未录完的需求，支持直接录入开发单号',
    '集团需求导出：筛选非敏捷需求，按页勾选后一键导出 Excel',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')

add_paragraph_cn('本系统为可分发版本，已完成以下安全与性能优化：')
opts = [
    'Web 端与插件端数据库连接池化，提升并发性能',
    '邮箱密码 AES-256-GCM 加密存储（Web 端 email-config.json 与 Chrome 插件 storage）',
    'Session 登录认证，替代浏览器基础认证',
    'SQL 注入防护',
    '插件本地 SMTP 中继服务器增加 CORS 来源校验，仅允许 Chrome 扩展调用敏感接口',
]
for o in opts:
    doc.add_paragraph(o, style='List Bullet')

# ==================== 2. 环境要求 ====================
add_heading_cn('2. 环境要求', level=1)
reqs = [
    '操作系统：Windows 10/11、Windows Server 2016+',
    'Node.js：v18 或更高版本',
    'MySQL：5.7 或更高版本（已存在目标数据库）',
    'Python：3.8+（用于数据库初始化脚本）',
    'Chrome 浏览器：用于安装需求提取kimi版插件',
    '本地 SMTP 中继服务器：如需邮件催办/插件发件功能，需运行 local-smtp-server（默认端口 2525）',
]
for r in reqs:
    doc.add_paragraph(r, style='List Bullet')

# ==================== 3. 项目结构 ====================
add_heading_cn('3. 项目结构说明', level=1)
add_paragraph_cn('项目主要目录及文件说明如下：')
add_code_block('''xqemail/
├── email-manager/              # 主应用目录
│   ├── config/                 # 配置加载器
│   ├── public/                 # 前端静态资源（HTML/CSS/JS）
│   │   ├── group.html          # 集团需求页面
│   │   ├── js/group.js         # 集团需求页面脚本
│   │   ├── todo.html           # 我的待办页面
│   │   ├── js/todo.js          # 我的待办页面脚本
│   │   └── lib/                # Excel 导出库
│   ├── routes/                 # API 路由
│   ├── utils/                  # 工具函数（认证、数据库、加密等）
│   ├── .env                    # 环境变量（需手动配置）
│   ├── .env.example            # 环境变量示例
│   ├── package.json            # Node.js 依赖配置
│   ├── server.js               # 服务端入口
│   ├── start-server.bat        # Windows 一键启动脚本
│   └── create_shortcut.py      # 创建桌面快捷方式
├── requirement-email-extractor/# Chrome 插件目录
│   ├── background/             # 后台 Service Worker
│   ├── content/                # 页面内容提取脚本
│   ├── icons/                  # 插件图标
│   ├── lib/                    # Excel 处理库
│   ├── local-smtp-server/      # 插件本地 SMTP 中继服务器
│   │   ├── routes/             # 服务端路由
│   │   ├── utils/              # 数据库、IMAP 工具
│   │   ├── server.js           # 中继服务器入口
│   │   ├── start-server.bat    # 可见窗口启动
│   │   └── start-silent.vbs    # 静默后台启动
│   ├── popup/                  # 插件弹窗页面
│   ├── manifest.json           # 插件配置
│   └── 安装说明.md              # 插件快速说明
├── database/                   # 数据库脚本
│   ├── add_columns.py          # 初始化 sent_emails 字段
│   └── check_table.py          # 查看表结构和数据
├── README.md                   # 项目说明
└── start-all.bat               # 一键启动 Web 服务 + SMTP 中继''')

# ==================== 4. 部署步骤 ====================
add_heading_cn('4. 部署步骤', level=1)

add_heading_cn('4.1 解压项目', level=2)
add_paragraph_cn('将分发包解压到目标目录，例如：')
add_code_block('C:\\WorkBuddy\\2026-06-18-09-00-23-distributable')

add_heading_cn('4.2 安装 Node.js 依赖', level=2)
add_paragraph_cn('打开命令提示符或 PowerShell，进入 email-manager 目录：')
add_code_block('cd C:\\WorkBuddy\\2026-06-18-09-00-23-distributable\\email-manager')
add_paragraph_cn('执行以下命令安装依赖：')
add_code_block('npm install')
add_paragraph_cn('安装成功后，目录下会生成 node_modules 文件夹。')

add_heading_cn('4.3 配置环境变量', level=2)
add_paragraph_cn('复制环境变量示例文件：')
add_code_block('copy .env.example .env')
add_paragraph_cn('使用文本编辑器打开 .env 文件，根据实际情况填写以下关键配置：')
add_code_block('''# 服务器端口
PORT=3000

# 登录账号密码（生产环境必须修改）
AUTH_USER=admin
AUTH_PASS=changeme

# MySQL 数据库配置
DB_HOST=127.0.0.1
DB_PORT=3306
DB_USER=root
DB_PASSWORD=你的真实数据库密码
DB_NAME=aicoding

# SMTP 发件邮箱配置
SMTP_HOST=smtp.example.com
SMTP_PORT=465
SMTP_SECURE=true
SMTP_USER=your_email@example.com
SMTP_PASS=your_email_password

# 配置文件加密密钥（建议 32 位以上随机字符串）
CONFIG_ENCRYPTION_KEY=your-32-char-encryption-key-here

# Session 密钥（建议长随机字符串）
SESSION_SECRET=your-long-random-session-secret-here''')
add_note('AUTH_USER、AUTH_PASS、CONFIG_ENCRYPTION_KEY、SESSION_SECRET 在生产环境中必须修改为强密码，切勿使用默认值。')

add_heading_cn('4.4 初始化数据库', level=2)
add_paragraph_cn('进入 database 目录，执行字段初始化脚本：')
add_code_block('''cd ..\\database
python add_columns.py''')
add_paragraph_cn('脚本会自动检查并添加以下字段：')
add_code_block('''workload        DECIMAL(10,2)   # 工作量（人天）
is_involved     TINYINT(1)      # 是否涉及开发（0=否，1=是）''')

add_heading_cn('4.5 启动服务', level=2)
add_paragraph_cn('本项目需要同时运行 Web 服务（端口 3000）和插件本地 SMTP 中继服务（端口 2525）。')
add_paragraph_cn('方式一（推荐，一键启动所有服务）：')
add_code_block('双击根目录下的 start-all.bat')
add_paragraph_cn('该脚本会自动打开两个窗口，分别运行需求催办 Web 服务和本地 SMTP 中继服务。')
add_paragraph_cn('方式二：分别启动')
add_code_block('''# 1. 启动 Web 服务
双击 email-manager\start-server.bat

# 2. 启动插件本地 SMTP 中继服务
双击 requirement-email-extractor\local-smtp-server\start-silent.vbs''')
add_paragraph_cn('方式三：命令行启动')
add_code_block('''cd email-manager
npm start''')
add_paragraph_cn('启动成功后，控制台会显示：')
add_code_block('''Web 服务运行在 http://0.0.0.0:3000
SMTP 中继服务运行在 http://127.0.0.1:2525''')

add_heading_cn('4.6 验证 Web 部署', level=2)
add_paragraph_cn('打开浏览器，访问以下地址：')
add_code_block('http://localhost:3000')
add_paragraph_cn('如果看到登录页面，说明 Web 服务部署成功。')
add_screenshot('01_login.png', '系统登录页面')

add_heading_cn('4.7 安装 Chrome 插件', level=2)
add_paragraph_cn('完成 Web 服务部署后，安装需求提取kimi版插件：')
plugin_steps = [
    '打开 Chrome 浏览器，地址栏输入 chrome://extensions/',
    '开启右上角“开发者模式”',
    '点击“加载已解压的扩展程序”',
    '选择 requirement-email-extractor 文件夹',
    '安装成功后，Chrome 工具栏会出现插件图标',
]
for i, s in enumerate(plugin_steps, 1):
    doc.add_paragraph(f'{i}. {s}')
add_screenshot('06_plugin_icon.png', 'Chrome 插件图标')
add_note('插件不依赖特定扩展 ID，任何人加载后均可使用。')

add_heading_cn('4.8 启动插件本地 SMTP 中继服务器', level=2)
add_paragraph_cn('插件发送邮件需要通过本地 SMTP 中继服务器。如果已经使用 start-all.bat 一键启动，则无需单独启动本服务。')
add_paragraph_cn('方式一（推荐，后台静默运行）：')
add_code_block('双击 requirement-email-extractor\\local-smtp-server\\start-silent.vbs')
add_paragraph_cn('方式二（可见窗口，方便查看日志）：')
add_code_block('双击 requirement-email-extractor\\local-smtp-server\\start-server.bat')
add_paragraph_cn('启动后，任务栏托盘不会有图标，但可在系统任务管理器中看到 node.exe 进程。服务器监听 127.0.0.1:2525，仅本机可访问。')
add_note('如果启动时提示端口被占用，请打开任务管理器结束 node.exe 进程后重试，或在命令行执行 netstat -ano | findstr :2525 查找占用进程。')

# ==================== 5. 操作手册 ====================
doc.add_page_break()
add_heading_cn('5. 日常使用操作手册', level=1)

add_heading_cn('5.1 登录系统', level=2)
add_paragraph_cn('在浏览器地址栏输入 http://localhost:3000，进入登录页面。')
add_paragraph_cn('输入用户名和密码，点击“登录”按钮。默认账号为：')
add_code_block('用户名：admin\n密码：changeme')
add_paragraph_cn('登录成功后自动跳转到需求清单首页。')

add_heading_cn('5.2 查看需求清单', level=2)
add_paragraph_cn('首页（需求清单）展示所有需求的分组列表，包含以下信息：')
cols = ['需求ID', '提出时间', '需求名称', '提单人', '涉及系统', '责任人', '工作量', '涉及开发', '开发单号']
doc.add_paragraph('、'.join(cols))
add_screenshot('02_index.png', '需求清单首页')
add_paragraph_cn('页面顶部会显示工作量初评提示卡片：自动统计未完成工作量初评的需求数量，并引导前往“催办提醒”页面进行催办，或直接在需求清单页录入初评工作量。')
add_paragraph_cn('页面顶部还显示统计卡片：总需求数、待填写工作量、已完成评估。')
add_paragraph_cn('涉及系统与责任人过多时，默认最多显示 2 个，超出部分以 +N 气泡展示，点击气泡可查看全部。')
add_paragraph_cn('责任人中不涉及开发的会以斜体显示，并标注“（不涉及开发）”；整体行仅展示涉及开发的责任人和系统。')
add_paragraph_cn('点击需求名称可查看详情弹窗。')
add_screenshot('05_detail.png', '需求详情弹窗')

add_heading_cn('5.3 编辑工作量', level=2)
steps = [
    '点击需求左侧的 ▶ 图标展开子行',
    '在“工作量”列输入人天数（支持 0.5 递增）',
    '在“涉及开发”列选择“是”或“否”',
    '如需填写开发单号，在“开发单号”列输入',
    '点击该行右侧的“保存”按钮可单独保存当前系统',
    '也可修改多行后点击顶部“保存修改”按钮进行批量保存',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')
add_note('填写大于 0 的工作量会自动将“涉及开发”设为“是”；取消涉及开发会自动清空工作量。')

add_heading_cn('5.4 发送邮件催办', level=2)
add_paragraph_cn('点击顶部导航“催办提醒”进入催办页面。')
add_screenshot('03_reminder.png', '需求催办页面')
steps = [
    '页面列出所有有待填写工作量的责任人',
    '点击责任人右侧的“邮件催办”按钮',
    '在确认弹窗中点击“确认发送”',
    '系统会自动从 sa_info 表获取责任人邮箱并发送催办邮件',
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {s}')
add_note('发送邮件前需确保：1）SMTP 邮箱配置正确；2）sa_info 表已维护责任人邮箱；3）本地 SMTP 中继服务器已启动。')

add_heading_cn('5.5 发送微信催办', level=2)
add_paragraph_cn('在催办页面点击“微信催办”按钮，系统会：')
wechat_steps = [
    '自动生成催办文案',
    '将文案复制到剪贴板',
    '弹出弹窗显示完整文案',
]
for s in wechat_steps:
    doc.add_paragraph(s, style='List Bullet')
add_paragraph_cn('用户可直接粘贴到微信发送。')

add_heading_cn('5.6 查看催办记录', level=2)
add_paragraph_cn('在需求清单首页点击顶部“催办记录”按钮，可查看最近 20 条催办记录，包括：')
record_cols = ['需求编号', '需求名称', '收件人', '状态', '发送时间']
doc.add_paragraph('、'.join(record_cols))

add_heading_cn('5.7 配置 Web 端邮箱', level=2)
add_paragraph_cn('在需求清单首页点击顶部“邮箱设置”按钮，打开邮箱配置弹窗。')
add_screenshot('04_email_config.png', '邮箱设置弹窗')
add_paragraph_cn('填写以下信息后点击“保存”：')
email_fields = [
    'SMTP 服务器、端口、是否 SSL/TLS',
    '发件人邮箱、密码/授权码',
    'IMAP 服务器、端口（可选）',
    '邮件签名、默认抄送（可选）',
]
for f in email_fields:
    doc.add_paragraph(f, style='List Bullet')
add_paragraph_cn('保存后可点击“测试连接”验证 SMTP 配置是否正确。')
add_note('邮箱密码保存后会自动加密存储到 email-config.json 文件中。')

add_heading_cn('5.8 退出登录', level=2)
add_paragraph_cn('点击顶部导航栏右侧的“退出登录”链接，即可安全退出系统并返回登录页。')

add_heading_cn('5.9 我的待办', level=2)
add_paragraph_cn('点击顶部导航“我的待办”进入待办页面。系统会自动筛选出满足以下条件的需求：')
todo_conditions = [
    '该需求至少有一个涉及开发的系统',
    '所有涉及开发的系统工作量均已评估（大于 0）',
    '至少有一个涉及开发的系统尚未录入开发单号',
]
for c in todo_conditions:
    doc.add_paragraph(c, style='List Bullet')
add_screenshot('11_todo.png', '我的待办页面')
add_paragraph_cn('页面顶部汇总当前待办需求数量，提醒及时转开发单并录入开发单号。')
add_paragraph_cn('待办清单的表格格式与需求清单页保持一致，可展开子行，在“开发单号”列输入单号后点击“保存”按钮。当该需求所有涉及开发的开发单号都录入完成后，该需求会自动从待办页面消失。')

add_heading_cn('5.10 集团需求导出', level=2)
add_paragraph_cn('点击顶部导航“集团需求”进入集团运营工单页面。')
group_steps = [
    '页面自动过滤需求ID中不包含“敏捷需求”的运营工单',
    '列表按工单发布时间从大到小排序，每页最多展示 10 条',
    '在需要导出的工单前勾选复选框，也可点击表头复选框全选当前页',
    '点击右上角“一键导出”按钮，系统自动生成 Excel 文件并下载',
]
for i, s in enumerate(group_steps, 1):
    doc.add_paragraph(f'{i}. {s}')
add_screenshot('10_group.png', '集团需求页面')
add_paragraph_cn('导出文件的表头会自动替换为：运营工单ID、工单发布时间、运营工单名称、发布人、开发单号或说明；文件名为“集团运营工单YYYYMMDDhhmmss.xlsx”。')
add_note('导出时只包含已勾选的工单，未勾选的数据不会写入 Excel。')

add_heading_cn('5.11 使用 Chrome 插件提取需求信息', level=2)
add_paragraph_cn('插件安装并配置完成后，可在任意需求/工单页面提取信息并发送邮件：')
plugin_usage_steps = [
    '打开目标需求页面（如 TAPD、Jira、内部工单系统等）',
    '点击 Chrome 工具栏的插件图标',
    '切换到“提取信息”标签，点击“提取需求信息”',
    '系统自动尝试识别需求编号、名称、提出人、背景、描述等字段',
    '如自动识别失败，可点击“点选”按钮，在页面上逐一点选字段值',
    '确认提取结果后，勾选收件人并点击“直接发送邮件”',
]
for i, s in enumerate(plugin_usage_steps, 1):
    doc.add_paragraph(f'{i}. {s}')
add_screenshot('07_plugin_extract.png', '插件提取信息界面')
add_paragraph_cn('发送成功后，插件会自动将需求信息写入 Web 端使用的 sent_emails 表，可在 Web 端“催办记录”中查看。')
add_note('发送前请确保：1）本地 SMTP 中继服务器已启动；2）插件邮箱设置中的 SMTP 信息正确；3）已选择至少一个收件人。')

add_heading_cn('5.12 插件收件人管理', level=2)
add_paragraph_cn('插件支持两种收件人来源：')
doc.add_paragraph('本地联系人：通过 Excel 导入或手动添加，仅保存在当前 Chrome 用户配置中', style='List Bullet')
doc.add_paragraph('数据库联系人：当插件数据库设置正确时，自动从 sa_info 表加载，并与 Web 端催办系统共享', style='List Bullet')
add_paragraph_cn('操作步骤：')
contact_steps = [
    '切换到插件的“收件人管理”标签',
    '点击“导入 Excel 文件”，选择包含“姓名、系统、邮箱”列的 .xlsx 文件',
    '或手动填写姓名、系统、邮箱后点击“添加”',
    '如使用数据库来源，点击“从数据库刷新”可同步 sa_info 表的最新数据',
]
for i, s in enumerate(contact_steps, 1):
    doc.add_paragraph(f'{i}. {s}')
add_screenshot('08_plugin_contacts.png', '插件收件人管理界面')

add_heading_cn('5.13 插件邮箱与数据库设置', level=2)
add_paragraph_cn('切换到插件的“邮箱设置”标签，按页面提示填写：')
plugin_email_fields = [
    'SMTP 服务器地址、端口、是否 SSL/TLS',
    '邮箱账号和授权码/密码',
    'IMAP 服务器地址（可选，用于存档到“已发送”文件夹）',
    '默认抄送邮箱、邮件签名',
    'MySQL 数据库主机、端口、库名、用户名、密码、表名',
]
for f in plugin_email_fields:
    doc.add_paragraph(f, style='List Bullet')
add_paragraph_cn('填写完成后，点击“测试发送”验证 SMTP，点击“测试连接”验证数据库。')
add_paragraph_cn('保存后，SMTP 密码和数据库密码都会先加密再写入 Chrome 本地存储，避免明文泄露。')
add_screenshot('09_plugin_settings.png', '插件邮箱设置界面')

# ==================== 6. 常见问题 ====================
doc.add_page_break()
add_heading_cn('6. 常见问题排查', level=1)

qa = [
    ('Q1：浏览器提示“无法访问此网站”',
     '请检查服务器是否已启动，以及 3000 端口是否被其他程序占用。'),
    ('Q2：登录后页面空白或数据加载失败',
     '请检查 .env 中的 DB_PASSWORD 是否正确，以及 MySQL 服务是否运行。'),
    ('Q3：发送邮件催办失败，提示中继服务器错误',
     '请确认本地 SMTP 中继服务器（127.0.0.1:2525）已启动，且 .env 中 SMTP 配置正确。'),
    ('Q4：邮箱配置保存后测试连接失败',
     '请确认 SMTP 服务器地址、端口、邮箱密码正确。部分邮箱需要使用授权码而非登录密码。'),
    ('Q5：如何修改默认登录密码',
     '编辑 email-manager/.env 文件，修改 AUTH_USER 和 AUTH_PASS，然后重启服务器。'),
    ('Q6：插件提示“本地邮件服务器未启动”',
     '请确认已双击 start-silent.vbs 或 start-server.bat 启动中继服务器。'),
    ('Q7：插件提取不到需求信息',
     '尝试使用“点选”模式手动指定字段；或在“提取规则”标签中配置更精确的 CSS 选择器。'),
    ('Q8：插件发送邮件后 Web 端看不到记录',
     '请检查插件的 MySQL 数据库设置是否与 Web 端一致，并确认 sent_emails 表名正确。'),
]
for q, a in qa:
    add_paragraph_cn(q, bold=True)
    add_paragraph_cn(a)
    doc.add_paragraph()

# ==================== 7. 安全建议 ====================
add_heading_cn('7. 安全建议', level=1)
safety = [
    '部署到生产环境前，务必修改默认账号密码和各项密钥',
    '使用 HTTPS 反向代理（如 Nginx）对外提供 Web 服务',
    '定期备份 MySQL 数据库',
    '不要将 .env、email-config.json 及插件源码中的敏感信息提交到版本控制',
    '限制 MySQL 数据库访问权限，避免使用 root 账号',
    '为邮箱账号开启两步验证，并使用专用授权码',
    '插件本地 SMTP 中继服务器仅监听 127.0.0.1，不要修改 HOST 配置暴露到公网',
    '离开电脑时锁定屏幕，防止他人直接使用已登录的 Chrome 插件',
]
for s in safety:
    doc.add_paragraph(s, style='List Bullet')

# 保存文档
doc.save(OUTPUT_DOC)
print(f'文档已生成：{OUTPUT_DOC}')
